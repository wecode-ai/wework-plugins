#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""Deterministic DOCX creation, inspection, editing, rendering, and validation."""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import zipfile
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def emit(payload: object, destination: str | None = None) -> None:
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    if destination:
        path = Path(destination)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(f"{text}\n", encoding="utf-8")
    else:
        print(text)


def read_spec(path: str) -> dict[str, object]:
    value = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise TypeError("document specification must be a JSON object")
    return value


def configure_document(document: DocumentType, spec: dict[str, object]) -> None:
    properties = document.core_properties
    properties.title = str(spec.get("title", ""))
    properties.author = str(spec.get("author", ""))
    properties.subject = str(spec.get("subject", ""))
    properties.keywords = str(spec.get("keywords", ""))
    normal = document.styles["Normal"]
    normal.font.name = str(spec.get("font", "Aptos"))
    normal.font.size = Pt(float(spec.get("fontSize", 11)))
    for section in document.sections:
        margins = spec.get("marginsInches", {})
        if isinstance(margins, dict):
            section.top_margin = Inches(float(margins.get("top", 0.8)))
            section.bottom_margin = Inches(float(margins.get("bottom", 0.8)))
            section.left_margin = Inches(float(margins.get("left", 0.9)))
            section.right_margin = Inches(float(margins.get("right", 0.9)))


def add_table(document: DocumentType, value: dict[str, object]) -> None:
    columns = value.get("columns", [])
    rows = value.get("rows", [])
    if not isinstance(columns, list) or not isinstance(rows, list):
        raise TypeError("table columns and rows must be arrays")
    width = len(columns) or max(
        (len(row) for row in rows if isinstance(row, list)),
        default=1,
    )
    table = document.add_table(rows=1 if columns else 0, cols=width)
    table.style = str(value.get("style", "Table Grid"))
    if columns:
        for index, cell_value in enumerate(columns):
            cell = table.rows[0].cells[index]
            cell.text = str(cell_value)
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for row_value in rows:
        if not isinstance(row_value, list):
            raise TypeError("every table row must be an array")
        cells = table.add_row().cells
        for index, cell_value in enumerate(row_value[:width]):
            cells[index].text = "" if cell_value is None else str(cell_value)


def add_content(document: DocumentType, items: object) -> None:
    if not isinstance(items, list):
        raise TypeError("content must be an array")
    for item in items:
        if not isinstance(item, dict):
            raise TypeError("every content item must be an object")
        if "heading" in item:
            document.add_heading(
                str(item["heading"]),
                level=int(item.get("level", 1)),
            )
        elif "paragraph" in item:
            paragraph = document.add_paragraph(
                str(item["paragraph"]),
                style=str(item.get("style", "Normal")),
            )
            alignment = str(item.get("alignment", "left")).lower()
            paragraph.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        elif "bullets" in item:
            values = item["bullets"]
            if not isinstance(values, list):
                raise ValueError("bullets must be an array")
            for value in values:
                document.add_paragraph(str(value), style="List Bullet")
        elif "numbered" in item:
            values = item["numbered"]
            if not isinstance(values, list):
                raise ValueError("numbered must be an array")
            for value in values:
                document.add_paragraph(str(value), style="List Number")
        elif "table" in item:
            table = item["table"]
            if not isinstance(table, dict):
                raise ValueError("table must be an object")
            add_table(document, table)
        elif "image" in item:
            image = Path(str(item["image"])).expanduser()
            document.add_picture(
                str(image),
                width=Inches(float(item.get("widthInches", 6.0))),
            )
        elif item.get("pageBreak") is True:
            document.add_page_break()
        else:
            raise ValueError(f"unsupported content item: {sorted(item)}")


def command_create(args: argparse.Namespace) -> None:
    spec = read_spec(args.spec)
    document = Document()
    configure_document(document, spec)
    title = str(spec.get("title", "")).strip()
    if title and spec.get("includeTitle", True):
        document.add_heading(title, level=0)
    add_content(document, spec.get("content", []))
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    emit({"status": "created", "path": str(output.resolve())})


def paragraph_record(index: int, paragraph: object) -> dict[str, object]:
    return {
        "index": index,
        "style": getattr(getattr(paragraph, "style", None), "name", None),
        "text": getattr(paragraph, "text", ""),
    }


def command_inspect(args: argparse.Namespace) -> None:
    source = Path(args.input)
    document = Document(source)
    payload = {
        "path": str(source.resolve()),
        "properties": {
            "title": document.core_properties.title or "",
            "author": document.core_properties.author or "",
            "subject": document.core_properties.subject or "",
            "keywords": document.core_properties.keywords or "",
        },
        "paragraphs": [
            paragraph_record(index, paragraph)
            for index, paragraph in enumerate(document.paragraphs)
        ],
        "tables": [
            {
                "index": table_index,
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
            }
            for table_index, table in enumerate(document.tables)
        ],
        "sections": [
            {
                "widthInches": round(section.page_width.inches, 3),
                "heightInches": round(section.page_height.inches, 3),
                "orientation": str(section.orientation),
            }
            for section in document.sections
        ],
    }
    emit(payload, args.output)


def all_paragraphs(document: DocumentType) -> Iterable[object]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_in_paragraph(paragraph: object, old: str, new: str) -> int:
    runs = list(getattr(paragraph, "runs", []))
    text = "".join(run.text for run in runs)
    count = text.count(old)
    if count == 0:
        return 0
    replacement = text.replace(old, new)
    if runs:
        runs[0].text = replacement
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.text = replacement
    return count


def command_replace(args: argparse.Namespace) -> None:
    document = Document(args.input)
    replacements = sum(
        replace_in_paragraph(paragraph, args.old, args.new)
        for paragraph in all_paragraphs(document)
    )
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    emit(
        {
            "status": "updated",
            "path": str(output.resolve()),
            "replacements": replacements,
        }
    )


def dependency_roots() -> list[Path]:
    configured = os.environ.get("CODEX_WORKSPACE_DEPENDENCIES", "").strip()
    if not configured:
        return []
    root = Path(configured).expanduser()
    return [root, root / "dependencies"]


def find_binary(name: str) -> str | None:
    resolved = shutil.which(name)
    if resolved:
        return resolved
    for root in dependency_roots():
        for candidate in (
            root / "bin" / "override" / name,
            root / "bin" / "fallback" / name,
            root / "native" / "poppler" / "bin" / name,
            root / "native" / "poppler" / "poppler" / "bin" / name,
        ):
            if candidate.is_file():
                return str(candidate)
    return None


def libreoffice_environment() -> dict[str, str]:
    environment = os.environ.copy()
    if environment.get("FONTCONFIG_FILE"):
        return environment
    for root in dependency_roots():
        configs = sorted(
            root.glob("native/libreoffice-headless/**/Resources/fontconfig/fonts.conf")
        )
        if configs:
            environment["FONTCONFIG_FILE"] = str(configs[0])
            environment["FONTCONFIG_PATH"] = str(configs[0].parent)
            break
    return environment


def command_render(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    office = find_binary("soffice") or find_binary("libreoffice")
    if not office:
        raise RuntimeError("LibreOffice is required to render DOCX files")
    subprocess.run(
        [
            office,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(source),
        ],
        check=True,
        capture_output=True,
        text=True,
        env=libreoffice_environment(),
    )
    pdf_path = output_dir / f"{source.stem}.pdf"
    if not pdf_path.is_file():
        raise RuntimeError("LibreOffice did not produce the expected PDF")
    images: list[str] = []
    if args.images:
        pdftoppm = find_binary("pdftoppm")
        if not pdftoppm:
            raise RuntimeError("pdftoppm is required when --images is requested")
        prefix = output_dir / source.stem
        subprocess.run(
            [pdftoppm, "-png", "-r", str(args.dpi), str(pdf_path), str(prefix)],
            check=True,
        )
        images = [str(path) for path in sorted(output_dir.glob(f"{source.stem}-*.png"))]
    emit({"pdf": str(pdf_path), "images": images})


def command_validate(args: argparse.Namespace) -> None:
    source = Path(args.input)
    issues: list[str] = []
    try:
        with zipfile.ZipFile(source) as archive:
            corrupt = archive.testzip()
            if corrupt:
                issues.append(f"corrupt ZIP entry: {corrupt}")
            names = set(archive.namelist())
            for required in ("[Content_Types].xml", "word/document.xml"):
                if required not in names:
                    issues.append(f"missing DOCX part: {required}")
        document = Document(source)
        for index, table in enumerate(document.tables):
            widths = {len(row.cells) for row in table.rows}
            if len(widths) > 1:
                issues.append(f"table {index} has inconsistent row widths")
        if not document.sections:
            issues.append("document has no sections")
    except Exception as error:  # noqa: BLE001 - validation reports format failures
        issues.append(str(error))
    emit({"valid": not issues, "issues": issues, "path": str(source.resolve())})
    if issues:
        raise SystemExit(1)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    create = subparsers.add_parser(
        "create", help="Create DOCX from a JSON specification"
    )
    create.add_argument("--spec", required=True)
    create.add_argument("--output", required=True)
    create.set_defaults(handler=command_create)

    inspect = subparsers.add_parser(
        "inspect", help="Inspect document content and structure"
    )
    inspect.add_argument("--input", required=True)
    inspect.add_argument("--output")
    inspect.set_defaults(handler=command_inspect)

    replace = subparsers.add_parser(
        "replace", help="Replace text throughout a document"
    )
    replace.add_argument("--input", required=True)
    replace.add_argument("--old", required=True)
    replace.add_argument("--new", required=True)
    replace.add_argument("--output", required=True)
    replace.set_defaults(handler=command_replace)

    render = subparsers.add_parser(
        "render", help="Render DOCX to PDF and optional PNG pages"
    )
    render.add_argument("--input", required=True)
    render.add_argument("--output-dir", required=True)
    render.add_argument("--images", action="store_true")
    render.add_argument("--dpi", type=int, default=144)
    render.set_defaults(handler=command_render)

    validate = subparsers.add_parser("validate", help="Validate DOCX package integrity")
    validate.add_argument("--input", required=True)
    validate.set_defaults(handler=command_validate)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    try:
        args.handler(args)
    except (OSError, ValueError, RuntimeError, subprocess.CalledProcessError) as error:
        print(f"documents: {error}", file=sys.stderr)
        raise SystemExit(2) from error


if __name__ == "__main__":
    main()
